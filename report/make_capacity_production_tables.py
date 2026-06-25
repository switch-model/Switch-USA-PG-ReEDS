#!/usr/bin/env python3
"""Create capacity and production summary tables for 2030 CE scenarios.

Run from the switch directory:

    python iterate/make_capacity_production_tables.py

By default this reads out/2030/ce/{high_fossil,clean} and writes one long CSV
with scenario and region columns. Region, scenario, and generator technology
groupings are read from make_load_balance_graphs.py without importing it.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

import numpy as np
import pandas as pd

# use python-docx package to generate tables in .docx document
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reporting_info import (
    rto_groups,
    region_info,
    scen_names,
    gen_tech_names,
    summary_out_dir,
    ce_out_dir,
    ce_in_dir,
    historical_year,
    model_year,
)

EE_LABEL = "Energy Efficiency"
DR_LABEL = "Demand Response"

METRIC_COLUMNS = [
    "existing_capacity_mw",
    "capacity_added_mw",
    "capacity_retired_mw",
    "capacity_online_mw",
    "production_gwh",
]

# soft hyphen for use in headers (will be converted to optional hyphen when
# creating the docx)
h_ = "\u00ad"

num_fmt = lambda x: "" if abs(x) < 0.1 else f"{x:,.1f}".replace("-", "−")

column_defs = [
    ("technology", "Technology", None, 1.5),
    ("existing_capacity_gw", "2025 capacity (GW)", num_fmt, 0.9),
    ("capacity_added_gw", "Capacity added (GW)", num_fmt, 0.9),
    ("capacity_retired_gw", "Capacity retired (GW)", num_fmt, 0.9),
    ("capacity_online_gw", "2030 capacity (GW)", num_fmt, 0.9),
    ("production_twh", f"Produc{h_}tion (TWh)", num_fmt, 0.9),
]

tech_order = [
    "Coal",
    "Gas CCGT",
    "Gas CT",
    "Gas Peaker",
    "Nuclear",
    "Convent. Hydro",
    "Onshore Wind",
    "Offshore Wind",
    "Large Solar",
    "Dist. Solar",
    "Geothermal",
    "Pumped Hydro",
    "Batteries",
    "Other",
    "U.S. Imports",
]
# these adjustments relative to make_load_balance_graphs.py
remap_tech_names = {
    "Batteries": "Batteries",
    "Other_peaker": "Gas Peaker",
    "Geothermal": "Geothermal",
    "Small Hydroelectric": "Convent. Hydro",
    "Conventional Hydroelectric": "Convent. Hydro",
    "Hydroelectric Pumped Storage": "Pumped Hydro",
    "Utility-Scale Battery Storage_Lithium Ion_Advanced": "Batteries",
    "Imports_base_base": "U.S. Imports",
    "load_growth": "omit",
}


def read_csv(path: Path, **kwargs) -> pd.DataFrame:
    return pd.read_csv(path, na_values=".", **kwargs)


def load_groupings():
    zone_groups = {
        group: region_info.loc[region_info["rto"].isin(rto_list), "region"].to_list()
        for group, rto_list in rto_groups.items()
    }

    gen_tech_names.update(remap_tech_names)
    technology_order = tech_order + [EE_LABEL, DR_LABEL]

    missing = set(gen_tech_names.values()) - {"omit"} - set(technology_order)
    if missing:
        raise ValueError(f"Technologies {missing} are missing from `tech_order` list.")

    return zone_groups, gen_tech_names, technology_order


def load_model_config(out_dir: Path) -> dict:
    config_file = out_dir / "model_config.json"
    if not config_file.exists():
        return {"options": {}}
    return json.loads(config_file.read_text())


def input_aliases(config: dict) -> dict[str, str]:
    aliases = {}
    for item in config.get("options", {}).get("input_aliases", []) or []:
        if "=" not in item:
            continue
        target, source = item.split("=", 1)
        aliases[target] = source
    return aliases


def scenario_input_path(config: dict, file_name: str) -> Path:
    options = config.get("options", {})
    file_name = input_aliases(config).get(file_name, file_name)
    return ce_in_dir / file_name


def read_scenario_input(config: dict, file_name: str, **kwargs) -> pd.DataFrame:
    return read_csv(scenario_input_path(config, file_name), **kwargs)


def standardize_project_col(df: pd.DataFrame) -> pd.DataFrame:
    if "GENERATION_PROJECT" in df.columns:
        return df
    if "generation_project" in df.columns:
        return df.rename(columns={"generation_project": "GENERATION_PROJECT"})
    raise KeyError("Expected generation_project or GENERATION_PROJECT column.")


def standardize_period_col(df: pd.DataFrame) -> pd.DataFrame:
    if "PERIOD" in df.columns:
        return df
    if "period" in df.columns:
        return df.rename(columns={"period": "PERIOD"})
    raise KeyError("Expected period or PERIOD column.")


def add_technology_labels(
    df: pd.DataFrame, gen_tech_names: dict[str, str]
) -> pd.DataFrame:
    missing = sorted(set(df["gen_tech"].dropna()) - set(gen_tech_names))
    if missing:
        missing_list = ", ".join(missing)
        raise ValueError(
            "Unexpected gen_tech values not mapped in make_load_balance_graphs.py: "
            f"{missing_list}"
        )
    df = df.copy()
    df["technology"] = df["gen_tech"].map(gen_tech_names)
    return df


def numeric(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce").fillna(0.0)


def target_period_info(
    periods: pd.DataFrame, target_period: int
) -> tuple[dict, int, int]:
    periods = periods.rename(columns={"INVESTMENT_PERIOD": "PERIOD"})
    periods["PERIOD"] = periods["PERIOD"].astype(int)
    periods = periods.set_index("PERIOD")

    if target_period in periods.index:
        period_start = int(periods.loc[target_period, "period_start"])
        period_end = int(periods.loc[target_period, "period_end"])
    else:
        period_start = target_period
        period_end = target_period

    period_start_by_period = periods["period_start"].astype(float).to_dict()
    return period_start_by_period, period_start, period_end


def online_in_year(df: pd.DataFrame, year: int) -> pd.Series:
    build_year = numeric(df["build_year"])
    max_age = numeric(df["gen_max_age"]).replace(0, np.inf)
    return (build_year <= year) & (year <= build_year + max_age)


def online_in_period(
    df: pd.DataFrame,
    target_period: int,
    period_start_by_period: dict,
    period_start: int,
    period_end: int,
    retire_time: str,
) -> pd.Series:
    build_year = numeric(df["build_year"])
    max_age = numeric(df["gen_max_age"]).replace(0, np.inf)
    online = build_year.map(period_start_by_period).fillna(build_year)
    retirement = online + max_age
    built_in_target_period = build_year == target_period

    if retire_time == "late":
        can_run = (online <= period_start) & (period_start < retirement)
    elif retire_time == "mid":
        mid_period = period_start + 0.5 * (period_end - period_start + 1)
        can_run = (online <= mid_period) & (mid_period <= retirement)
    elif retire_time == "early":
        can_run = (online < period_end) & (period_end <= retirement)
    elif retire_time == "early_early":
        can_run = (online < period_end) & (period_end < retirement)
    else:
        raise ValueError(f"Unexpected retire_time option {retire_time!r}")

    return built_in_target_period | can_run


def aggregate_by_region(
    df: pd.DataFrame,
    zone_col: str,
    zone_groups: dict[str, list[str]],
    scenario: str,
) -> pd.DataFrame:
    rows = []
    for region, zones in zone_groups.items():
        sub = df.loc[df[zone_col].isin(zones)]
        if sub.empty:
            continue
        grouped = sub.groupby("technology", as_index=False)[METRIC_COLUMNS].sum()
        grouped.insert(0, "scenario", scenario)
        grouped.insert(0, "region", region)
        rows.append(grouped)

    if rows:
        return pd.concat(rows, ignore_index=True)
    return pd.DataFrame(columns=["scenario", "region", "technology"] + METRIC_COLUMNS)


def zero_metric_frame(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    for col in METRIC_COLUMNS:
        if col not in df.columns:
            df[col] = 0.0
    return df


def concat_metric_frames(frames: list[pd.DataFrame]) -> pd.DataFrame:
    frames = [frame for frame in frames if not frame.empty]
    if not frames:
        return pd.DataFrame(
            columns=["scenario", "region", "technology"] + METRIC_COLUMNS
        )
    return pd.concat(frames, ignore_index=True)


def build_capacity_metrics(
    scenario: str,
    out_dir: Path,
    config: dict,
    gen_tech_names: dict[str, str],
    zone_groups: dict[str, list[str]],
    base_year: int,
    target_period: int,
) -> pd.DataFrame:
    periods = read_scenario_input(config, "periods.csv")
    period_start_by_period, period_start, period_end = target_period_info(
        periods, target_period
    )
    retire_time = config.get("options", {}).get("retire_time", "late")

    gen_info = standardize_project_col(read_scenario_input(config, "gen_info.csv"))
    gen_info = gen_info[
        [
            "GENERATION_PROJECT",
            "gen_tech",
            "gen_load_zone",
            "gen_energy_source",
            "gen_max_age",
        ]
    ]

    build = standardize_project_col(read_csv(out_dir / "gen_build.csv"))
    build["BuildGen"] = numeric(build["BuildGen"])
    build["build_year"] = numeric(build["build_year"])
    build = build.merge(
        gen_info[["GENERATION_PROJECT", "gen_max_age"]],
        on="GENERATION_PROJECT",
        how="left",
        validate="many_to_one",
    )
    if build["gen_max_age"].isna().any():
        missing = sorted(
            build.loc[build["gen_max_age"].isna(), "GENERATION_PROJECT"].unique()
        )
        raise ValueError(f"Missing gen_max_age for generation projects: {missing}")
    build = add_technology_labels(build, gen_tech_names)

    build["online_in_base_year"] = online_in_year(build, base_year)
    build["online_in_target_period"] = online_in_period(
        build,
        target_period,
        period_start_by_period,
        period_start,
        period_end,
        retire_time,
    )
    build["existing_capacity_mw"] = np.where(
        build["online_in_base_year"], build["BuildGen"], 0.0
    )
    build["capacity_added_mw"] = np.where(
        build["build_year"].between(base_year + 1, target_period),
        build["BuildGen"],
        0.0,
    )
    build["capacity_retired_mw"] = np.where(
        build["online_in_base_year"] & ~build["online_in_target_period"],
        build["BuildGen"],
        0.0,
    )

    build_metrics = aggregate_by_region(
        zero_metric_frame(build),
        "gen_load_zone",
        zone_groups,
        scenario,
    )

    suspend_metrics = switch_retirement_metrics(
        scenario,
        out_dir,
        gen_info,
        gen_tech_names,
        zone_groups,
        base_year,
        target_period,
    )

    return concat_metric_frames([build_metrics, suspend_metrics])


def switch_retirement_metrics(
    scenario: str,
    out_dir: Path,
    gen_info: pd.DataFrame,
    gen_tech_names: dict[str, str],
    zone_groups: dict[str, list[str]],
    base_year: int,
    target_period: int,
) -> pd.DataFrame:
    suspend_file = out_dir / "SuspendGen.csv"
    if suspend_file.exists():
        suspend = read_csv(suspend_file).rename(
            columns={
                "GEN_BLD_SUSPEND_YRS_1": "GENERATION_PROJECT",
                "GEN_BLD_SUSPEND_YRS_2": "build_year",
                "GEN_BLD_SUSPEND_YRS_3": "PERIOD",
            }
        )
        suspend["PERIOD"] = numeric(suspend["PERIOD"]).astype(int)
        suspend["build_year"] = numeric(suspend["build_year"])
        suspend["SuspendGen"] = numeric(suspend["SuspendGen"])
        suspend = suspend.loc[
            (suspend["PERIOD"] == target_period)
            & (suspend["build_year"] <= base_year)
            & (suspend["SuspendGen"] != 0)
        ]
        if suspend.empty:
            return pd.DataFrame(
                columns=["scenario", "region", "technology"] + METRIC_COLUMNS
            )
        suspend = suspend.merge(
            gen_info,
            on="GENERATION_PROJECT",
            how="left",
            validate="many_to_one",
        )
        if suspend["gen_max_age"].isna().any():
            missing = sorted(
                suspend.loc[
                    suspend["gen_max_age"].isna(), "GENERATION_PROJECT"
                ].unique()
            )
            raise ValueError(f"Missing gen_info rows for suspended projects: {missing}")
        suspend = add_technology_labels(suspend, gen_tech_names)
        suspend["capacity_retired_mw"] = suspend["SuspendGen"]
        return aggregate_by_region(
            zero_metric_frame(suspend),
            "gen_load_zone",
            zone_groups,
            scenario,
        )

    cap = standardize_project_col(
        standardize_period_col(read_csv(out_dir / "gen_cap.csv"))
    )
    cap = cap.loc[numeric(cap["PERIOD"]).astype(int) == target_period].copy()
    cap["capacity_retired_mw"] = numeric(cap.get("SuspendGen_total", 0.0))
    cap = cap.loc[cap["capacity_retired_mw"] != 0]
    if cap.empty:
        return pd.DataFrame(
            columns=["scenario", "region", "technology"] + METRIC_COLUMNS
        )
    cap = add_technology_labels(cap, gen_tech_names)
    return aggregate_by_region(
        zero_metric_frame(cap), "gen_load_zone", zone_groups, scenario
    )


def online_capacity_metrics(
    scenario: str,
    out_dir: Path,
    gen_tech_names: dict[str, str],
    zone_groups: dict[str, list[str]],
    target_period: int,
) -> pd.DataFrame:
    cap = standardize_project_col(
        standardize_period_col(read_csv(out_dir / "gen_cap.csv"))
    )
    cap = cap.loc[numeric(cap["PERIOD"]).astype(int) == target_period].copy()
    cap["capacity_online_mw"] = numeric(cap["GenCapacity"])
    cap = cap.loc[cap["capacity_online_mw"] != 0]
    if cap.empty:
        return pd.DataFrame(
            columns=["scenario", "region", "technology"] + METRIC_COLUMNS
        )
    cap = add_technology_labels(cap, gen_tech_names)
    return aggregate_by_region(
        zero_metric_frame(cap), "gen_load_zone", zone_groups, scenario
    )


def production_metrics(
    scenario: str,
    out_dir: Path,
    gen_tech_names: dict[str, str],
    zone_groups: dict[str, list[str]],
    target_period: int,
) -> pd.DataFrame:
    annual_summary = out_dir / "dispatch_zonal_annual_summary.csv"
    if not annual_summary.exists():
        annual_summary = out_dir / "dispatch_gen_annual_summary.csv"

    dispatch = standardize_period_col(read_csv(annual_summary))
    dispatch = dispatch.loc[
        numeric(dispatch["PERIOD"]).astype(int) == target_period
    ].copy()
    dispatch["production_gwh"] = numeric(dispatch["Energy_GWh_typical_yr"])
    dispatch = add_technology_labels(dispatch, gen_tech_names)
    return aggregate_by_region(
        zero_metric_frame(dispatch),
        "gen_load_zone",
        zone_groups,
        scenario,
    )


def investment_metrics(
    scenario: str,
    out_dir: Path,
    zone_groups: dict[str, list[str]],
    target_period: int,
) -> pd.DataFrame:
    rows = []

    ee = read_optional_investment(out_dir / "energy_efficiency_investment.csv")
    if ee is not None:
        ee = standardize_period_col(ee)
        ee = ee.loc[numeric(ee["PERIOD"]).astype(int) == target_period].copy()
        ee["technology"] = EE_LABEL
        ee["capacity_added_mw"] = numeric(ee.get("DeployEEMW", 0.0))
        ee["capacity_online_mw"] = numeric(ee.get("DeployEEMW", 0.0))
        ee["production_gwh"] = numeric(ee.get("DeployEEGWh_per_year", 0.0))
        rows.append(
            aggregate_by_region(
                zero_metric_frame(ee),
                "load_zone",
                zone_groups,
                scenario,
            )
        )

    dr = read_optional_investment(out_dir / "demand_response_investment.csv")
    if dr is not None:
        dr = standardize_period_col(dr)
        dr = dr.loc[numeric(dr["PERIOD"]).astype(int) == target_period].copy()
        dr["technology"] = DR_LABEL
        dr["capacity_added_mw"] = numeric(dr.get("DeployDRMW", 0.0))
        dr["capacity_online_mw"] = numeric(dr.get("DeployDRMW", 0.0))
        rows.append(
            aggregate_by_region(
                zero_metric_frame(dr),
                "load_zone",
                zone_groups,
                scenario,
            )
        )

    if rows:
        return concat_metric_frames(rows)
    return pd.DataFrame(columns=["scenario", "region", "technology"] + METRIC_COLUMNS)


def read_optional_investment(path: Path) -> pd.DataFrame | None:
    if path.exists():
        return read_csv(path)
    else:
        # print(f"{path} not found; treating as zero.")
        return None


def scenario_metrics(
    scenario: str,
    ce_out_dir: Path,
    gen_tech_names: dict[str, str],
    zone_groups: dict[str, list[str]],
    base_year: int,
    target_period: int,
) -> pd.DataFrame:
    out_dir = ce_out_dir / scenario
    if not out_dir.is_dir():
        raise FileNotFoundError(f"Scenario output directory not found: {out_dir}")

    config = load_model_config(out_dir)
    frames = [
        build_capacity_metrics(
            scenario,
            out_dir,
            config,
            gen_tech_names,
            zone_groups,
            base_year,
            target_period,
        ),
        online_capacity_metrics(
            scenario,
            out_dir,
            gen_tech_names,
            zone_groups,
            target_period,
        ),
        production_metrics(
            scenario,
            out_dir,
            gen_tech_names,
            zone_groups,
            target_period,
        ),
        investment_metrics(scenario, out_dir, zone_groups, target_period),
    ]
    return concat_metric_frames(frames)


def finalize_tables(
    metrics: pd.DataFrame,
    scenarios: list[str],
    zone_groups: dict[str, list[str]],
    technology_order: list[str],
) -> pd.DataFrame:
    grouped = (
        metrics.groupby(["scenario", "region", "technology"], as_index=False)[
            METRIC_COLUMNS
        ]
        .sum()
        .set_index(["scenario", "region", "technology"])
    )

    full_index = pd.MultiIndex.from_product(
        [scenarios, list(zone_groups), technology_order],
        names=["scenario", "region", "technology"],
    )
    grouped = grouped.reindex(full_index, fill_value=0.0).reset_index()

    # change to bigger units
    grouped[METRIC_COLUMNS] *= 0.001
    result = grouped.rename(
        columns={
            c: c.replace("_gwh", "_twh").replace("_mw", "_gw") for c in METRIC_COLUMNS
        }
    )

    return result


######
# docx table generation
def set_cell_shading(cell, fill):
    """Set cell background color, e.g. fill='D9EAF7'."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def format_value(x, fmt=None):
    if pd.isna(x):
        return ""
    if fmt is not None:
        return fmt(x)
    if isinstance(x, float):
        return f"{x:,.1f}"
    if isinstance(x, int):
        return f"{x:,}"
    return str(x)


def set_paragraph_text(paragraph, text):
    """
    Replace paragraph contents with text, converting Unicode soft hyphens
    to Word-native <w:softHyphen/> elements.
    """
    paragraph.clear()
    parts = text.split(h_)
    for i, part in enumerate(parts):
        if part:
            paragraph.add_run(part)
        if i < len(parts) - 1:
            run = paragraph.add_run()
            run._r.append(OxmlElement("w:softHyphen"))


def add_dataframe_table(
    doc,
    df,
    title,
    font_name="Aptos",
    font_size=9,
    table_style="Table Grid",
    column_widths=None,
    number_formats=None,
):
    """
    Add a pandas DataFrame as a formatted Word table.

    column_widths: dict mapping column name to width in inches
    number_formats: dict mapping column name to function, e.g.
        {"Cost": lambda x: f"${x:,.0f}", "Share": lambda x: f"{x:.1%}"}
    """
    number_formats = number_formats or {}

    # Title
    p = doc.add_paragraph()
    p.style = doc.styles["Caption"]
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    # run.font.name = font_name
    # run.font.size = Pt(10)
    # run.bold = True

    # Table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = table_style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header row
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        set_paragraph_text(hdr_cells[j].paragraphs[0], str(col))
        set_cell_shading(hdr_cells[j], "ECECEC")
        hdr_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Body rows
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            fmt = number_formats.get(col)
            cells[j].text = format_value(val, fmt)

    # Formatting
    numeric_cols = set(df.select_dtypes(include="number").columns)

    for row_idx, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            col = df.columns[j]

            if column_widths and col in column_widths:
                set_cell_width(cell, column_widths[col])

            for p in cell.paragraphs:
                if row_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif col in numeric_cols:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                for run in p.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    if row_idx == 0:
                        run.bold = True

                # take out any line and paragraph spacing; rely on Table Grid style only
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0

                # prevent page breaks between rows, but allow after last row
                if row_idx + 1 < len(table.rows):
                    p.paragraph_format.keep_with_next = True

    doc.add_paragraph()  # spacing after table
    return table


def create_docx(file, result):
    rename = {r[0]: r[1] for r in column_defs}
    widths = {r[1]: r[3] for r in column_defs}
    formats = {r[1]: r[2] for r in column_defs if r[2] is not None}

    doc = Document()

    # specify custom line and paragraph spacing for Normal style, so the table
    # rows can explicitly override it (to 1/0), and that will carry through a
    # copy-paste into a different document with custom line and paragraph
    # spacing
    normal = doc.styles["Normal"]
    normal.paragraph_format.space_before = Pt(2)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.2

    group_cols = ["region", "scenario"]
    for (region, scenario), group in result.groupby(group_cols, sort=False):
        add_dataframe_table(
            doc,
            group.drop(columns=group_cols).rename(columns=rename),
            title=(
                f"Electricity generation capacity and production, "
                f"{region}, {scenario.replace('_', ' ')}"
            ),
            font_name="Montserrat",
            font_size=10,
            table_style="Table Grid",
            column_widths=widths,
            number_formats=formats,
        )

    Path(file).parent.mkdir(parents=True, exist_ok=True)
    doc.save(file)
    print(f"Wrote {file}.")


def main() -> None:
    output_file = summary_out_dir / "capacity_production_tables.docx"

    zone_groups, gen_tech_names, technology_order = load_groupings()

    metrics = pd.concat(
        [
            scenario_metrics(
                scenario,
                ce_out_dir,
                gen_tech_names,
                zone_groups,
                historical_year,
                model_year,
            )
            for scenario in scen_names
        ],
        ignore_index=True,
    )
    result = finalize_tables(metrics, scen_names, zone_groups, technology_order)
    # imports aren't really capacity (and exact availability is unknown)
    result.loc[
        result["technology"] == "U.S. Imports",
        ["capacity_added_gw", "capacity_online_gw"],
    ] = 0
    # we don't know existing dist PV or additions, so just show final amount
    result.loc[
        result["technology"] == "Dist. Solar",
        ["existing_capacity_gw", "capacity_added_gw"],
    ] = 0
    result = result.sort_values(["region", "scenario"], kind="stable")
    create_docx(output_file, result)

    # output_file.parent.mkdir(parents=True, exist_ok=True)
    # result.to_csv(output_file, index=False)
    # print(f"Wrote {len(result)} rows to {output_file}")


if __name__ == "__main__":
    main()
